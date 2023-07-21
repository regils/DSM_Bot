import sqlite3
import pandas
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import os
import datetime

month_names = {
    1: 'январь',
    2: 'февраль',
    3: 'март',
    4: 'апрель',
    5: 'май',
    6: 'июнь',
    7: 'июль',
    8: 'август',
    9: 'сентябрь',
    10: 'октябрь',
    11: 'ноябрь',
    12: 'декабрь'
}

async def users_start():
    global conn, cursor
    conn = sqlite3.connect('users_list.db')
    cursor = conn.cursor()

    cursor.execute("CREATE TABLE IF NOT EXISTS users (user_id TEXT PRIMARY KEY, user_name TEXT, tasks INTEGER)")
    conn.commit()

async def answers_start():
    cursor.execute("CREATE TABLE IF NOT EXISTS answers (user_id TEXT, user_name TEXT, q1 TEXT, q2 TEXT, q3 TEXT)")
    conn.commit()

async def prev_answers_start():
    cursor.execute("CREATE TABLE IF NOT EXISTS previous_answers (user_id TEXT, user_name TEXT, q1 TEXT, q2 TEXT, q3 TEXT)")
    conn.commit()

async def clear_answers():
    cursor.execute("DELETE FROM answers")
    conn.commit()

async def clear_prev_answers():
    cursor.execute("DELETE FROM previous_answers")
    conn.commit()

async def copy_answers():
    cursor.execute("INSERT INTO previous_answers SELECT * FROM answers")
    conn.commit()

async def create_profile(user_id, user_name):
    user = cursor.execute("SELECT 1 FROM users WHERE user_id = ?", (user_id,)).fetchone()
    if not user:
        cursor.execute("INSERT INTO users VALUES(?, ?, ?)", (user_id, user_name, 0))
        conn.commit()
    else:
        cursor.execute("UPDATE users SET user_name = ? WHERE user_id = ?", (user_name, user_id,))
        conn.commit()

async def update_name(user_id, user_name):
    cursor.execute("UPDATE answers SET user_name = ? WHERE user_id = ?", (user_name, user_id,))
    conn.commit()

async def reset_tasks():
    cursor.execute("UPDATE users SET tasks = 0")
    conn.commit()

async def reset_tasks_for_user(user_id):
    cursor.execute("UPDATE users SET tasks = 0 WHERE user_id = ?", (user_id,))
    conn.commit()

async def add_task(user_id):
    cursor.execute("UPDATE users SET tasks = tasks + 1 WHERE user_id = ?", (user_id,))
    conn.commit()

async def get_tasks1(user_id):
    user = cursor.execute("SELECT tasks FROM users WHERE user_id = ?", (user_id,)).fetchone()
    tasks = user[0]

    return tasks

async def get_tasks(user_id):
    tasks = cursor.execute("SELECT * FROM answers where user_id = ?", (user_id,)).fetchall()
    return len(tasks)

async def update_tasks(user_id):
    cursor.execute("UPDATE users SET tasks = ? WHERE user_id = ?", (await get_tasks(user_id), user_id))
    conn.commit()

async def delete_tasks(user_id):
    cursor.execute("DELETE FROM answers WHERE user_id = ?", (user_id,))
    conn.commit()
    await update_tasks(user_id)

async def delete_answer_from_user(user_id, q1, q2, q3):
    cursor.execute("DELETE FROM answers WHERE user_id = ? AND q1 = ? AND q2 = ? AND q3 = ?", (user_id, q1, q2, q3))
    conn.commit()
    await update_tasks(user_id)

async def get_answers(user_id):
    answer_list = cursor.execute("SELECT q1, q2, q3 FROM answers WHERE user_id = ?", (user_id,)).fetchall()
    if answer_list:
        return answer_list
    else:
        return None
    
async def get_prev_answers(user_id):
    answer_list = cursor.execute("SELECT q1, q2, q3 FROM previous_answers WHERE user_id = ?", (user_id,)).fetchall()
    if answer_list:
        return answer_list
    else:
        return None

async def select_answer(user_id, q1):
    answer_list = cursor.execute("SELECT * FROM answers WHERE user_id = ? AND q1 = ?", (user_id, q1)).fetchall()
    if answer_list:
        return answer_list
    else:
        return None
    
async def edit_q1(user_id, q1, q1_new):
    cursor.execute("UPDATE answers SET q1 = ? WHERE user_id = ? AND q1 = ?", (q1_new, user_id, q1))
    conn.commit()

async def edit_q2(user_id, q1, q2_new):
    cursor.execute("UPDATE answers SET q2 = ? WHERE user_id = ? AND q1 = ?", (q2_new, user_id, q1))
    conn.commit()

async def edit_q3(user_id, q1, q3_new):
    cursor.execute("UPDATE answers SET q3 = ? WHERE user_id = ? AND q1 = ?", (q3_new, user_id, q1))
    conn.commit()

async def get_name(user_id):
    user = cursor.execute("SELECT user_name FROM users WHERE user_id = ?", (user_id,)).fetchone()
    
    if user is not None:
        user_name = user[0]
        return user_name
    else:
        return None

async def get_users():
    users = cursor.execute("SELECT * FROM users").fetchall()
    return users

async def create_answer(user_id, q1, q2, q3):
    user_name = await get_name(user_id=user_id)
    cursor.execute("INSERT INTO answers (user_id, user_name, q1, q2, q3) "
                "VALUES (?, ?, ?, ?, ?)", (user_id, user_name, q1, q2, q3))
    conn.commit()

async def sort_answers():
    cursor.execute("SELECT * FROM answers ORDER BY user_name")
    conn.commit()

async def delete_report():
    if os.path.exists('report.docx'):
        os.remove('report.docx')

async def create_report():
    cursor.execute("SELECT q1, user_name, q2, q3 FROM answers ORDER BY user_name")
    rows = cursor.fetchall()
    df = pandas.DataFrame(rows, columns=['Задание', 'Исполнитель', 'Срок исполнения', 'Результат выполнения'])
    doc = Document()

    current_date = datetime.datetime.now()
    current_month = current_date.month
    current_year = current_date.year
    current_month_ru = month_names[current_month]

    doc.add_paragraph(f"Отчет о работе отдела разработки автоматизированных систем (отдел № 2), сектора информационно-измерительных систем (сектор № 21) СКБ «Меридиан» за {current_month_ru} {current_year} г.")

    table = doc.add_table(df.shape[0] + 1, df.shape[1] + 3)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    table.cell(0, 0).text = '№ п/п'

    for i, column_name in enumerate(df.columns):
        table.cell(0, i + 1).text = column_name

    current_executor = None
    executor_count = 0

    current_executor_value = None
    current_executor_value_row = None

    for i, row in enumerate(df.itertuples(), start=1):
        executor = row.Исполнитель
        if executor != current_executor:
            executor_count += 1
            current_executor = executor
        table.cell(i, 0).text = str(executor_count)
        
        executor_value = row.Исполнитель
        if executor_value != current_executor_value:
            current_executor_value = executor_value
            current_executor_value_row = i
        else:
            table.cell(i, 2).text = ""
            table.cell(i, 2).merge(table.cell(current_executor_value_row, 2))
            table.cell(i, 0).text = ""
            table.cell(i, 0).merge(table.cell(current_executor_value_row, 0))
            table.cell(i, 5).merge(table.cell(current_executor_value_row, 5))
            table.cell(i, 6).merge(table.cell(current_executor_value_row, 6))
        
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j + 1).text = str(value)

    table.cell(0, df.shape[1] + 1).text = 'КТУ'
    table.cell(0, df.shape[1] + 2).text = 'Примечание'

    doc.save('report.docx')


